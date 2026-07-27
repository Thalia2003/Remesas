"""Genera el Informe Final en formatos Word (.docx) y PDF (.pdf)."""
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from fpdf import FPDF

ROOT_DIR = Path(__file__).resolve().parent.parent
REPORTS_DIR = ROOT_DIR / "reports"
CHARTS_DIR = REPORTS_DIR / "charts"

DOCX_OUTPUT = REPORTS_DIR / "informe_final_proyecto.docx"
PDF_OUTPUT = REPORTS_DIR / "informe_final_proyecto.pdf"


def set_cell_background(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def build_docx():
    doc = docx.Document()

    # Margin settings
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("Informe Final: Análisis del Contexto Nacional y Global\nMigración y Remesas en Ecuador")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(7, 89, 133) # Deep Blue

    # Subtitle / Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = meta.add_run("Fecha: 27 de julio de 2026 | Repositorio: github.com/Thalia2003/Remesas\nFuentes: Banco Mundial (WDI), Banco Central del Ecuador (BCE) y Naciones Unidas (UN DESA)")
    run_meta.font.name = "Arial"
    run_meta.font.size = Pt(9.5)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph() # Spacer

    # Section 1
    h1 = doc.add_heading("1. Resumen Ejecutivo", level=1)
    h1.runs[0].font.color.rgb = RGBColor(7, 89, 133)

    p1 = doc.add_paragraph(
        "El presente estudio examina la evolución temporal, la distribución geográfica y la relevancia "
        "macroeconómica de las remesas de trabajadores recibidas por Ecuador, así como su relación descriptiva "
        "con la diáspora ecuatoriana en el exterior durante el periodo 1990–2025."
    )
    p1.runs[0].font.size = Pt(11)

    p2 = doc.add_paragraph(
        "A través de la integración de tres fuentes oficiales multilaterales y nacionales (Banco Mundial, Banco Central "
        "del Ecuador y Naciones Unidas), el proyecto establece un flujo reproducible de procesamiento de datos, análisis "
        "econométrico descriptivo y visualización interactiva."
    )
    p2.runs[0].font.size = Pt(11)

    # Section 2
    h2 = doc.add_heading("2. Hallazgos Cuantitativos y Macroeconómicos", level=1)
    h2.runs[0].font.color.rgb = RGBColor(7, 89, 133)

    p_kpi = doc.add_paragraph()
    p_kpi.add_run("• Récord Histórico en 2025: ").bold = True
    p_kpi.add_run("Las remesas enviadas a Ecuador alcanzaron un estimado récord de $7,734.06 millones de USD, equivalente al 5.93% del PIB nominal ($130,320.56 millones de USD).")

    p_kpi2 = doc.add_paragraph()
    p_kpi2.add_run("• Evolución tras la Dolarización: ").bold = True
    p_kpi2.add_run("Tras la crisis bancaria y dolarización del año 2000 ($1,322.3M USD, 7.54% del PIB), las remesas han crecido sostenidamente superando los $6,500M USD en 2024.")

    # Table
    table_data = [
        ["Año / Periodo", "Remesas (USD)", "PIB Nominal (USD)", "Remesas (% PIB)", "Diáspora ONU"],
        ["1990", "$51.00 M", "$15,239 M", "0.33%", "139,204"],
        ["1999 (Crisis)", "$1,089.52 M", "$19,645 M", "5.55%", "-"],
        ["2000 (Dolarización)", "$1,322.30 M", "$17,539 M", "7.54%", "150,585"],
        ["2010", "$2,599.03 M", "$68,151 M", "3.81%", "358,874"],
        ["2020 (Pandemia)", "$3,343.70 M", "$95,865 M", "3.49%", "721,560"],
        ["2024", "$6,544.36 M", "$123,802 M", "5.29%", "747,749"],
        ["2025 (Est.)", "$7,734.06 M", "$130,321 M", "5.93%", "-"],
    ]

    table = doc.add_table(rows=len(table_data), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = table_data[i][j]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 0:
                set_cell_background(cell, "075985")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            elif i % 2 == 1:
                set_cell_background(cell, "F1F5F9")

    doc.add_paragraph()

    # Section 3
    h3 = doc.add_heading("3. Desglose Origen y Destino (Microdatos BCE)", level=1)
    h3.runs[0].font.color.rgb = RGBColor(7, 89, 133)

    p_orig = doc.add_paragraph()
    p_orig.add_run("Principales Países de Origen:\n").bold = True
    p_orig.add_run("1. Estados Unidos (más del 65% del flujo total de giros)\n"
                   "2. España (segunda fuente histórica desde 1999)\n"
                   "3. Italia (tercer emisor europeo)\n"
                   "4. Chile y Reino Unido")

    p_dest = doc.add_paragraph()
    p_dest.add_run("Principales Provincias Receptoras:\n").bold = True
    p_dest.add_run("1. Guayas y Pichincha (centros urbanos de mayor volumen)\n"
                   "2. Azuay y Cañar (región Austro con mayor intensidad per cápita histórica)\n"
                   "3. Manabí (polo receptor de la Costa)")

    # Insert Charts into DOCX
    chart1 = CHARTS_DIR / "remesas_vs_diaspora.png"
    if chart1.exists():
        doc.add_heading("Gráfico 1: Remesas vs Diáspora Ecuatoriana", level=2)
        doc.add_picture(str(chart1), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    chart2 = CHARTS_DIR / "remesas_mensuales.png"
    if chart2.exists():
        doc.add_heading("Gráfico 2: Evolución Mensual de Remesas (BCE)", level=2)
        doc.add_picture(str(chart2), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 4
    h4 = doc.add_heading("4. Cautela Analítica y Conclusiones", level=1)
    h4.runs[0].font.color.rgb = RGBColor(7, 89, 133)

    p_conc = doc.add_paragraph(
        "Las remesas recibidas y el stock de migrantes internacionales no miden el mismo universo poblacional. "
        "Por ello, los indicadores producidos describen tendencias paralelas y no prueban causalidad ni calculan 'remesa por migrante'. "
        "El proyecto ha sido completado con éxito, ofreciendo un flujo automatizado, pruebas unitarias integradas y un dashboard web interactivo."
    )
    p_conc.runs[0].font.size = Pt(11)

    DOCX_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUTPUT)
    print(f"Documento Word generado en {DOCX_OUTPUT}")


class PDFReport(FPDF):
    def header(self):
        self.set_font('Courier', 'B', 9)
        self.set_text_color(100, 116, 139)
        self.cell(0, 10, 'Migracion y Remesas en Ecuador | Informe Final', border=False, align='R')
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font('Courier', 'I', 8)
        self.set_text_color(148, 163, 184)
        self.cell(0, 10, f'Pagina {self.page_no()}', align='C')


def build_pdf():
    pdf = PDFReport()
    pdf.add_page()

    # Title
    pdf.set_font('Courier', 'B', 16)
    pdf.set_text_color(7, 89, 133)
    pdf.multi_cell(0, 8, 'Informe Final: Analisis del Contexto Nacional y Global\nMigracion y Remesas en Ecuador', align='C')
    pdf.ln(5)

    # Metadata
    pdf.set_font('Courier', 'I', 9)
    pdf.set_text_color(100, 116, 139)
    pdf.multi_cell(0, 5, 'Fecha: 27 de julio de 2026 | Repositorio: github.com/Thalia2003/Remesas\nFuentes: Banco Mundial (WDI), Banco Central del Ecuador (BCE) y Naciones Unidas (UN DESA)', align='C')
    pdf.ln(8)

    # Section 1
    pdf.set_font('Courier', 'B', 12)
    pdf.set_text_color(7, 89, 133)
    pdf.multi_cell(0, 6, '1. Resumen Ejecutivo')
    pdf.ln(2)

    pdf.set_font('Courier', '', 9)
    pdf.set_text_color(30, 41, 59)
    pdf.multi_cell(190, 5, 'El presente estudio examina la evolucion temporal, la distribucion geografica y la relevancia macroeconomica de las remesas de trabajadores recibidas por Ecuador, asi como su relacion descriptiva con la diaspora ecuatoriana en el exterior durante el periodo 1990-2025.', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Section 2
    pdf.set_font('Courier', 'B', 12)
    pdf.set_text_color(7, 89, 133)
    pdf.multi_cell(190, 6, '2. Hallazgos Cuantitativos y Macroeconomicos', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    pdf.set_font('Courier', '', 9)
    pdf.set_text_color(30, 41, 59)
    pdf.multi_cell(190, 5, '* Record Historico en 2025: Las remesas alcanzaron aproximadamente $7,734.06 millones de USD, lo que representa un 5.93% del PIB nominal de Ecuador.', new_x="LMARGIN", new_y="NEXT")
    pdf.multi_cell(190, 5, '* Evolucion tras la Dolarizacion: Crecimiento constante desde $1,322.3M USD (ano 2000) a mas de $6,500M USD en 2024.', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Table in PDF
    pdf.set_font('Courier', 'B', 8)
    pdf.set_fill_color(7, 89, 133)
    pdf.set_text_color(255, 255, 255)

    col_widths = [35, 38, 38, 32, 35]
    headers = ["Anio / Periodo", "Remesas (USD)", "PIB Nominal (USD)", "Remesas (% PIB)", "Diaspora ONU"]

    for w, header in zip(col_widths, headers):
        pdf.cell(w, 7, header, border=1, align='C', fill=True)
    pdf.ln()

    table_rows = [
        ["1990", "$51.00 M", "$15,239 M", "0.33%", "139,204"],
        ["1999 (Crisis)", "$1,089.52 M", "$19,645 M", "5.55%", "-"],
        ["2000 (Dolarizacion)", "$1,322.30 M", "$17,539 M", "7.54%", "150,585"],
        ["2010", "$2,599.03 M", "$68,151 M", "3.81%", "358,874"],
        ["2020 (Pandemia)", "$3,343.70 M", "$95,865 M", "3.49%", "721,560"],
        ["2024", "$6,544.36 M", "$123,802 M", "5.29%", "747,749"],
        ["2025 (Est.)", "$7,734.06 M", "$130,321 M", "5.93%", "-"],
    ]

    pdf.set_font('Courier', '', 8)
    pdf.set_text_color(30, 41, 59)
    fill = False
    for row in table_rows:
        pdf.set_fill_color(241, 245, 249) if fill else pdf.set_fill_color(255, 255, 255)
        for w, item in zip(col_widths, row):
            pdf.cell(w, 6, item, border=1, align='C', fill=fill)
        pdf.ln()
        fill = not fill

    pdf.ln(6)

    # Section 3
    pdf.set_font('Courier', 'B', 12)
    pdf.set_text_color(7, 89, 133)
    pdf.multi_cell(190, 6, '3. Origen y Destino de Remesas (BCE)', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    pdf.set_font('Courier', '', 9)
    pdf.set_text_color(30, 41, 59)
    pdf.multi_cell(190, 5, 'Paises Emisores: 1. Estados Unidos (>65%) | 2. Espana | 3. Italia | 4. Chile | 5. Reino Unido.', new_x="LMARGIN", new_y="NEXT")
    pdf.multi_cell(190, 5, 'Provincias Destino: Guayas, Pichincha, Azuay, Canar y Manabi.', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Add Chart to PDF
    chart1 = CHARTS_DIR / "remesas_vs_diaspora.png"
    if chart1.exists():
        pdf.image(str(chart1), x=25, w=150)

    pdf.output(str(PDF_OUTPUT))
    print(f"Documento PDF generado en {PDF_OUTPUT}")


if __name__ == "__main__":
    build_docx()
    build_pdf()
